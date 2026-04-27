from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional
from docx import Document
import re
from datetime import datetime
from pathlib import Path
import os

app = FastAPI(title="Analyst-X DOCX Export API")

REPORT_DIR = Path("/tmp/reports")
REPORT_DIR.mkdir(parents=True, exist_ok=True)

API_KEY = os.getenv("ANALYST_X_API_KEY")


class Section(BaseModel):
    number: int
    title: str
    type: str
    content: Optional[List[str]] = None
    columns: Optional[List[str]] = None
    rows: Optional[List[List[str]]] = None


class Metadata(BaseModel):
    company_name: str
    date: str
    time: str
    company_type: str
    research_goal: str
    language: str
    comparison_company: str
    confidence: str
    confidence_rationale: str


class Report(BaseModel):
    metadata: Metadata
    sections: List[Section]


class ExportRequest(BaseModel):
    file_name: Optional[str] = None
    report: Report


def check_api_key(request: Request):
    if not API_KEY:
        return
    provided_key = request.headers.get("x-api-key")
    if provided_key != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid API key")


def safe_filename(name: str) -> str:
    name = re.sub(r"[^\w\-.]+", "_", name, flags=re.UNICODE)
    return name.strip("_")


def clean_part(value: str) -> str:
    value = value or "Unknown"
    value = value.replace(" ", "_")
    return safe_filename(value)


def add_table(doc: Document, columns: List[str], rows: List[List[str]]):
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = str(col)

    for row in rows:
        cells = table.add_row().cells
        for i in range(len(columns)):
            cells[i].text = str(row[i]) if i < len(row) and row[i] is not None else ""


def validate_report(report: Report):
    section_numbers = {s.number for s in report.sections}
    missing = [i for i in range(1, 16) if i not in section_numbers]
    if missing:
        raise HTTPException(status_code=400, detail=f"Missing sections: {missing}")

    for section in report.sections:
        if section.type not in ["paragraphs", "table"]:
            raise HTTPException(status_code=400, detail=f"Invalid type in section {section.number}")

        if section.type == "paragraphs":
            if section.content is None or len(section.content) == 0:
                raise HTTPException(status_code=400, detail=f"Missing content in section {section.number}")

        if section.type == "table":
            if not section.columns:
                raise HTTPException(status_code=400, detail=f"Missing table columns in section {section.number}")
            if not section.rows or len(section.rows) == 0:
                raise HTTPException(status_code=400, detail=f"Table in section {section.number} must have at least one row")

            source_index = len(section.columns) - 1
            for row_i, row in enumerate(section.rows):
                if len(row) <= source_index or not str(row[source_index]).strip():
                    raise HTTPException(
                        status_code=400,
                        detail=f"Empty Source in section {section.number}, row {row_i + 1}"
                    )


@app.get("/")
def root():
    return {"status": "ok", "message": "Analyst-X DOCX Export API is running"}


@app.post("/generate-docx")
def generate_docx(payload: ExportRequest, request: Request):
    check_api_key(request)

    print("Incoming request for:", payload.report.metadata.company_name)

    validate_report(payload.report)
    meta = payload.report.metadata

    report_time = meta.time if meta.time and meta.time.strip() not in ["—", "-", ""] else datetime.now().strftime("%H-%M")
    report_time = report_time.replace(":", "-")

    if payload.file_name:
        file_name = payload.file_name
    else:
        company = clean_part(meta.company_name)
        goal = clean_part(meta.research_goal)
        file_name = f"{company}_{goal}_{meta.date}_{report_time}.docx"

    file_name = safe_filename(file_name)
    if not file_name.endswith(".docx"):
        file_name += ".docx"

    output_path = REPORT_DIR / file_name

    doc = Document()
    doc.add_heading(meta.company_name, level=1)
    doc.add_heading("Market & Competitive Intelligence Report", level=2)

    doc.add_paragraph(f"Date: {meta.date}")
    doc.add_paragraph(f"Time: {meta.time}")
    doc.add_paragraph(f"Company Type: {meta.company_type}")
    doc.add_paragraph(f"Research Goal: {meta.research_goal}")
    doc.add_paragraph(f"Language: {meta.language}")
    doc.add_paragraph(f"Comparison Company: {meta.comparison_company}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Confidence: {meta.confidence}")
    doc.add_paragraph(f"Confidence rationale: {meta.confidence_rationale}")

    for section in sorted(payload.report.sections, key=lambda s: s.number):
        doc.add_heading(f"{section.number}. {section.title}", level=2)

        if section.type == "table":
            add_table(doc, section.columns or [], section.rows or [])
        else:
            for paragraph in section.content or []:
                doc.add_paragraph(str(paragraph))

    doc.save(output_path)

    return {
        "status": "success",
        "file_name": file_name,
        "download_url": f"https://analyst-x-docx-api-production.up.railway.app/download/{file_name}"
    }


@app.get("/download/{file_name}")
def download_file(file_name: str):
    safe_name = safe_filename(file_name)
    file_path = REPORT_DIR / safe_name

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=safe_name
    )
